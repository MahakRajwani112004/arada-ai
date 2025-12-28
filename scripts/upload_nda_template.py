#!/usr/bin/env python3
"""Create and upload the NDA DOCX template to MinIO storage.

This script creates a professionally formatted NDA template document
with placeholders using {{PLACEHOLDER}} syntax, then uploads it to MinIO.

Usage:
    python scripts/upload_nda_template.py
"""

import asyncio
import sys
from pathlib import Path

# Add project root to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from src.storage.object_storage import get_storage
from src.config.logging import get_logger

logger = get_logger(__name__)

# Storage key for the NDA template
NDA_TEMPLATE_STORAGE_KEY = "templates/nda/mutual_nda_template.docx"


def create_nda_template() -> bytes:
    """Create a professionally formatted NDA template document."""
    doc = Document()

    # Set up document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('MUTUAL NON-DISCLOSURE AGREEMENT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.bold = True

    doc.add_paragraph()

    # Effective date paragraph
    para = doc.add_paragraph()
    para.add_run("This Mutual Non-Disclosure Agreement (the \"Agreement\") is made and entered into on ")
    effective_date_run = para.add_run("{{EFFECTIVE_DATE}}")
    effective_date_run.bold = True
    para.add_run(" (\"Effective Date\") by and between:")

    doc.add_paragraph()

    # First Party (Magure)
    first_party_heading = doc.add_paragraph()
    first_party_heading.add_run("FIRST PARTY:").bold = True

    first_party = doc.add_paragraph()
    first_party.add_run("Magure Softwares Private Limited").bold = True
    doc.add_paragraph("Address: MIG-10, Sector 2, Shankar Nagar, Raipur, Chhattisgarh, India - 492007")

    doc.add_paragraph()
    doc.add_paragraph("AND")
    doc.add_paragraph()

    # Second Party (placeholder)
    second_party_heading = doc.add_paragraph()
    second_party_heading.add_run("SECOND PARTY:").bold = True

    second_party = doc.add_paragraph()
    company_name_run = second_party.add_run("{{SECOND_PARTY_NAME}}")
    company_name_run.bold = True

    address_para = doc.add_paragraph()
    address_para.add_run("Address: ")
    address_run = address_para.add_run("{{SECOND_PARTY_ADDRESS}}")
    address_run.bold = True

    doc.add_paragraph()
    doc.add_paragraph("(each being referred to as a \"Party\" and together, the \"Parties\")")
    doc.add_paragraph()

    # Purpose section
    purpose_heading = doc.add_paragraph()
    purpose_heading.add_run("PURPOSE:").bold = True

    purpose_para = doc.add_paragraph()
    purpose_para.add_run("The Parties acknowledge and agree to keep the contents of the information and documentation obtained in respect of ")
    purpose_run = purpose_para.add_run("{{PURPOSE}}")
    purpose_run.bold = True
    purpose_para.add_run(" (\"Purpose\") confidential by signing this Non-Disclosure Agreement.")

    doc.add_paragraph()

    # Section 1 - Confidential Information
    section1 = doc.add_heading("1. CONFIDENTIAL INFORMATION", level=1)
    doc.add_paragraph(
        "\"Confidential Information\" shall mean all information shared orally or in writing including but not "
        "limited to financial, technical and non-technical information related to a party's business and its "
        "current, future and proposed products as well as services, research, development, design details and "
        "specifications, financial information, customer lists, business forecasts, sales information, marketing "
        "plans, and evaluation material."
    )

    # Section 2 - Confidentiality Obligations
    doc.add_heading("2. CONFIDENTIALITY OBLIGATIONS", level=1)
    doc.add_paragraph(
        "Recipient will not use any Confidential Information except to the extent necessary for the Purpose "
        "described above. Recipient will not disseminate or disclose any Confidential Information to any third "
        "party, except as expressly permitted in this Agreement. Recipient shall treat all of Discloser's "
        "Confidential Information with the same degree of care as Recipient accords to its own Confidential "
        "Information, but not less than reasonable care."
    )

    # Section 3 - Exclusions
    doc.add_heading("3. EXCLUSIONS", level=1)
    doc.add_paragraph("Recipient's obligations do not apply to information that:")
    doc.add_paragraph("(a) is publicly available through no fault of Recipient", style='List Bullet')
    doc.add_paragraph("(b) is rightfully in Recipient's possession free of any obligation of confidence", style='List Bullet')
    doc.add_paragraph("(c) is developed independently without use of Confidential Information", style='List Bullet')
    doc.add_paragraph("(d) is communicated by Discloser to an unaffiliated third party free of any obligation", style='List Bullet')

    # Section 4 - Term
    doc.add_heading("4. TERM", level=1)
    doc.add_paragraph(
        "This Agreement shall automatically terminate upon (i) completion of 2 (two) years from the Effective "
        "Date or (ii) by execution of a definitive Agreement between the Parties. Recipient's confidentiality "
        "obligations shall continue for 2 years post-termination."
    )

    # Section 5 - Governing Law
    doc.add_heading("5. GOVERNING LAW", level=1)
    doc.add_paragraph(
        "The laws of Republic of India govern all matters arising out of this Agreement. All disputes shall be "
        "subject to the exclusive jurisdiction of the courts situated in Raipur."
    )

    # Section 6 - Entire Agreement
    doc.add_heading("6. ENTIRE AGREEMENT", level=1)
    doc.add_paragraph(
        "This Agreement constitutes the final and exclusive agreement between the parties with respect to the "
        "treatment of Confidential Information disclosed hereunder."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature section
    sig_intro = doc.add_paragraph()
    sig_intro.add_run("IN WITNESS WHEREOF, ").bold = True
    sig_intro.add_run("the Parties have agreed to all the terms set forth in this Agreement.")

    doc.add_paragraph()
    doc.add_paragraph()

    # Magure signature block
    magure_sig = doc.add_paragraph()
    magure_sig.add_run("FOR AND ON BEHALF OF:").bold = True
    doc.add_paragraph()
    magure_company = doc.add_paragraph()
    magure_company.add_run("MAGURE SOFTWARES PRIVATE LIMITED").bold = True

    doc.add_paragraph()
    doc.add_paragraph("Signature: _______________________________")
    doc.add_paragraph("Name: Mohamed Oueida")
    doc.add_paragraph("Title: Director")

    date_para1 = doc.add_paragraph()
    date_para1.add_run("Date: ")
    date_run1 = date_para1.add_run("{{EFFECTIVE_DATE}}")
    date_run1.bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Second party signature block
    second_sig = doc.add_paragraph()
    second_sig.add_run("FOR AND ON BEHALF OF:").bold = True
    doc.add_paragraph()
    second_company = doc.add_paragraph()
    company_run2 = second_company.add_run("{{SECOND_PARTY_NAME}}")
    company_run2.bold = True

    doc.add_paragraph()
    doc.add_paragraph("Signature: _______________________________")

    name_para = doc.add_paragraph()
    name_para.add_run("Name: ")
    name_run = name_para.add_run("{{SIGNATORY_NAME}}")
    name_run.bold = True

    title_para = doc.add_paragraph()
    title_para.add_run("Title: ")
    title_run = title_para.add_run("{{SIGNATORY_TITLE}}")
    title_run.bold = True

    date_para2 = doc.add_paragraph()
    date_para2.add_run("Date: ")
    date_run2 = date_para2.add_run("{{EFFECTIVE_DATE}}")
    date_run2.bold = True

    # Save to bytes
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


async def upload_template():
    """Upload the NDA template to MinIO storage."""
    print("Creating NDA template document...")
    template_bytes = create_nda_template()
    print(f"Template created: {len(template_bytes)} bytes")

    print(f"Uploading to storage at: {NDA_TEMPLATE_STORAGE_KEY}")
    storage = get_storage()

    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    await storage.upload(NDA_TEMPLATE_STORAGE_KEY, template_bytes, content_type=content_type)

    print("Upload successful!")

    # Get a test URL to verify
    test_url = await storage.get_url(NDA_TEMPLATE_STORAGE_KEY, expires_seconds=3600)
    print(f"Template available at: {test_url}")

    return NDA_TEMPLATE_STORAGE_KEY


if __name__ == "__main__":
    storage_key = asyncio.run(upload_template())
    print(f"\nNDA Template Storage Key: {storage_key}")
    print("\nUse this storage key in the skill configuration.")
