"""DOCX Template Filler Tool - Fills placeholders in DOCX templates while preserving formatting.

This tool allows agents to:
1. Load an existing DOCX template from storage
2. Find and replace {{PLACEHOLDER}} markers in the document
3. Preserve all formatting, headers, footers, styles
4. Store the filled document in MinIO
5. Return a download URL
"""

import re
import uuid
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional

from src.config.logging import get_logger
from src.storage.object_storage import get_storage
from src.tools.base import BaseTool, ToolDefinition, ToolParameter, ToolResult

logger = get_logger(__name__)

GENERATED_FILES_PREFIX = "generated"


class DocxTemplateFiller(BaseTool):
    """
    Tool for filling DOCX templates while preserving document formatting.

    Unlike TemplateFillerTool which works with text, this tool:
    - Loads actual DOCX files from storage
    - Preserves headers, footers, styles, fonts, tables, etc.
    - Replaces placeholders in-place within the document structure
    """

    @property
    def definition(self) -> ToolDefinition:
        """Return DOCX template filler tool definition."""
        return ToolDefinition(
            name="fill_docx_template",
            description="""Fill an existing DOCX template file with data while preserving all formatting.

Use this when you need to:
- Fill a pre-formatted Word document template
- Preserve headers, footers, styles, fonts, and layout
- Replace {{PLACEHOLDER}} markers with actual values

The tool loads the template from storage, replaces placeholders, and returns a download URL.
Unlike generate_document or fill_template, this preserves the original document's formatting.""",
            parameters=[
                ToolParameter(
                    name="template_storage_key",
                    type="string",
                    description="The storage key/path of the DOCX template file (e.g., 'templates/nda_template.docx')",
                    required=True,
                ),
                ToolParameter(
                    name="values",
                    type="object",
                    description="Key-value pairs to fill in. Keys should match placeholder names without braces (e.g., {'COMPANY_NAME': 'Acme Corp', 'DATE': 'January 1, 2025'})",
                    required=True,
                ),
                ToolParameter(
                    name="output_filename",
                    type="string",
                    description="Filename for the generated document (without extension)",
                    required=True,
                ),
            ],
        )

    async def execute(self, **kwargs: Any) -> ToolResult:
        """Fill DOCX template and return download URL."""
        template_storage_key = kwargs.get("template_storage_key", "")
        values = kwargs.get("values", {})
        output_filename = kwargs.get("output_filename", "document")

        if not template_storage_key:
            return ToolResult(
                success=False,
                output=None,
                error="template_storage_key is required",
            )

        if not values:
            return ToolResult(
                success=False,
                output=None,
                error="values dictionary is required to fill the template",
            )

        try:
            from docx import Document
        except ImportError:
            return ToolResult(
                success=False,
                output=None,
                error="python-docx is required. Install with: pip install python-docx",
            )

        try:
            storage = get_storage()

            # Download the template from storage
            logger.info("downloading_template", storage_key=template_storage_key)
            template_bytes = await storage.download(template_storage_key)

            if not template_bytes:
                return ToolResult(
                    success=False,
                    output=None,
                    error=f"Template not found at: {template_storage_key}",
                )

            # Load the DOCX document
            import io
            doc = Document(io.BytesIO(template_bytes))

            # Track replacements for reporting
            replacements_made = {}
            unfilled_placeholders = set()

            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_in_paragraph(paragraph, values, replacements_made, unfilled_placeholders)

            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, values, replacements_made, unfilled_placeholders)

            # Replace placeholders in headers
            for section in doc.sections:
                for header in [section.header, section.first_page_header, section.even_page_header]:
                    if header:
                        for paragraph in header.paragraphs:
                            self._replace_in_paragraph(paragraph, values, replacements_made, unfilled_placeholders)
                        for table in header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        self._replace_in_paragraph(paragraph, values, replacements_made, unfilled_placeholders)

                # Replace placeholders in footers
                for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                    if footer:
                        for paragraph in footer.paragraphs:
                            self._replace_in_paragraph(paragraph, values, replacements_made, unfilled_placeholders)
                        for table in footer.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        self._replace_in_paragraph(paragraph, values, replacements_made, unfilled_placeholders)

            # Save the filled document to bytes
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            filled_bytes = output_buffer.read()

            # Generate unique file ID and upload
            file_id = f"doc_{uuid.uuid4().hex[:12]}"
            safe_filename = re.sub(r'[^\w\s-]', '', output_filename).strip().replace(' ', '_')
            filename = f"{safe_filename}.docx"
            storage_key = f"{GENERATED_FILES_PREFIX}/{file_id}/{filename}"

            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            await storage.upload(storage_key, filled_bytes, content_type=content_type)

            # Get presigned download URL (24 hour expiry)
            download_url = await storage.get_url(storage_key, expires_seconds=86400)

            logger.info(
                "docx_template_filled",
                file_id=file_id,
                filename=filename,
                template=template_storage_key,
                replacements=len(replacements_made),
                unfilled=len(unfilled_placeholders),
            )

            result = {
                "file_id": file_id,
                "filename": filename,
                "format": "docx",
                "size_bytes": len(filled_bytes),
                "download_url": download_url,
                "expires_at": (datetime.utcnow() + timedelta(hours=24)).isoformat(),
                "replacements_made": replacements_made,
                "message": f"Document '{filename}' generated successfully from template. Download link valid for 24 hours.",
            }

            if unfilled_placeholders:
                result["warnings"] = [f"Unfilled placeholders: {', '.join(sorted(unfilled_placeholders))}"]

            return ToolResult(
                success=True,
                output=result,
            )

        except Exception as e:
            logger.error("docx_template_fill_failed", error=str(e), template=template_storage_key)
            return ToolResult(
                success=False,
                output=None,
                error=f"Failed to fill DOCX template: {str(e)}",
            )

    def _replace_in_paragraph(
        self,
        paragraph,
        values: Dict[str, str],
        replacements_made: Dict[str, str],
        unfilled_placeholders: set,
    ) -> None:
        """Replace placeholders in a paragraph while preserving formatting.

        This handles the case where placeholders might be split across multiple runs
        (Word often splits text into multiple runs for formatting reasons).
        """
        # First, get the full text to find placeholders
        full_text = paragraph.text

        # Find all placeholders in this paragraph
        placeholder_pattern = r'\{\{(\w+)\}\}'
        matches = list(re.finditer(placeholder_pattern, full_text))

        if not matches:
            return

        # Track which placeholders we found
        for match in matches:
            placeholder_name = match.group(1)
            if placeholder_name in values:
                replacements_made[placeholder_name] = values[placeholder_name]
            else:
                unfilled_placeholders.add(placeholder_name)

        # For simple cases where each run contains complete placeholders
        for run in paragraph.runs:
            run_text = run.text
            for placeholder_name, value in values.items():
                placeholder = f"{{{{{placeholder_name}}}}}"
                if placeholder in run_text:
                    run.text = run_text.replace(placeholder, str(value))
                    run_text = run.text

        # Handle split placeholders (placeholder text split across multiple runs)
        # Rebuild paragraph text and check if any placeholders remain
        new_full_text = paragraph.text
        remaining_placeholders = re.findall(placeholder_pattern, new_full_text)

        if remaining_placeholders:
            # Some placeholders are split across runs - need to merge and re-split
            self._handle_split_placeholders(paragraph, values, replacements_made, unfilled_placeholders)

    def _handle_split_placeholders(
        self,
        paragraph,
        values: Dict[str, str],
        replacements_made: Dict[str, str],
        unfilled_placeholders: set,
    ) -> None:
        """Handle placeholders that are split across multiple runs.

        This is a more aggressive approach that consolidates runs when needed.
        """
        # Get full text and find remaining placeholders
        full_text = paragraph.text
        placeholder_pattern = r'\{\{(\w+)\}\}'

        # If there are placeholders, we need to handle them
        if not re.search(placeholder_pattern, full_text):
            return

        # Replace in the full text
        new_text = full_text
        for placeholder_name, value in values.items():
            placeholder = f"{{{{{placeholder_name}}}}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))
                replacements_made[placeholder_name] = value

        # If text changed, update the paragraph
        if new_text != full_text:
            # Clear all runs except first, put all text in first run
            if paragraph.runs:
                # Preserve formatting of first run
                first_run = paragraph.runs[0]

                # Clear all runs
                for run in paragraph.runs[1:]:
                    run.text = ""

                # Set text on first run
                first_run.text = new_text
