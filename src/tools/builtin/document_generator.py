"""Document Generator Tool - Creates downloadable documents from templates.

This tool allows agents to:
1. Fill in template placeholders with content
2. Generate documents in various formats (DOCX, PDF, TXT)
3. Store generated files in MinIO
4. Return download URLs to users
"""

import json
import re
import tempfile
import uuid
from datetime import datetime, timedelta
from typing import Any, Dict, Optional

from src.config.logging import get_logger
from src.storage.object_storage import get_storage
from src.tools.base import BaseTool, ToolDefinition, ToolParameter, ToolResult

logger = get_logger(__name__)

# Generated files bucket/prefix
GENERATED_FILES_PREFIX = "generated"


class DocumentGeneratorTool(BaseTool):
    """
    Tool for generating documents from templates or content.

    Agents can use this tool to:
    - Create documents from scratch with structured content
    - Fill in template placeholders
    - Output as DOCX, PDF, or TXT
    """

    @property
    def definition(self) -> ToolDefinition:
        """Return document generator tool definition."""
        return ToolDefinition(
            name="generate_document",
            description="""Generate a downloadable document. Use this when the user needs an actual file (not just text).

Supports:
- Creating documents from content (contracts, invoices, letters, reports)
- Filling templates with data
- Output formats: docx, pdf, txt

The tool returns a download URL that expires in 24 hours.""",
            parameters=[
                ToolParameter(
                    name="title",
                    type="string",
                    description="Document title/filename (without extension)",
                    required=True,
                ),
                ToolParameter(
                    name="content",
                    type="string",
                    description="The full document content. Use markdown formatting for structure.",
                    required=True,
                ),
                ToolParameter(
                    name="output_format",
                    type="string",
                    description="Output format: 'docx', 'pdf', or 'txt'",
                    required=False,
                    default="docx",
                    enum=["docx", "pdf", "txt"],
                ),
                ToolParameter(
                    name="metadata",
                    type="object",
                    description="Optional metadata (author, date, etc.)",
                    required=False,
                ),
            ],
        )

    async def execute(self, **kwargs: Any) -> ToolResult:
        """Generate a document and return download URL."""
        title = kwargs.get("title", "document")
        content = kwargs.get("content", "")
        output_format = kwargs.get("output_format", "docx").lower()
        metadata = kwargs.get("metadata", {})

        if not content:
            return ToolResult(
                success=False,
                output=None,
                error="Content is required to generate a document",
            )

        if output_format not in ["docx", "pdf", "txt"]:
            return ToolResult(
                success=False,
                output=None,
                error=f"Unsupported format: {output_format}. Use 'docx', 'pdf', or 'txt'",
            )

        try:
            # Generate unique file ID
            file_id = f"doc_{uuid.uuid4().hex[:12]}"
            safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
            filename = f"{safe_title}.{output_format}"
            storage_key = f"{GENERATED_FILES_PREFIX}/{file_id}/{filename}"

            # Generate document based on format
            if output_format == "txt":
                file_bytes = await self._generate_txt(content, metadata)
                content_type = "text/plain"
            elif output_format == "docx":
                file_bytes = await self._generate_docx(content, title, metadata)
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif output_format == "pdf":
                file_bytes = await self._generate_pdf(content, title, metadata)
                content_type = "application/pdf"

            # Upload to MinIO
            storage = get_storage()
            await storage.upload(storage_key, file_bytes, content_type=content_type)

            # Get presigned download URL (24 hour expiry)
            download_url = await storage.get_url(storage_key, expires_seconds=86400)

            logger.info(
                "document_generated",
                file_id=file_id,
                filename=filename,
                format=output_format,
                size_bytes=len(file_bytes),
            )

            return ToolResult(
                success=True,
                output={
                    "file_id": file_id,
                    "filename": filename,
                    "format": output_format,
                    "size_bytes": len(file_bytes),
                    "download_url": download_url,
                    "expires_at": (datetime.utcnow() + timedelta(hours=24)).isoformat(),
                    "message": f"Document '{filename}' generated successfully. Download link valid for 24 hours.",
                },
            )

        except Exception as e:
            logger.error("document_generation_failed", error=str(e), title=title)
            return ToolResult(
                success=False,
                output=None,
                error=f"Failed to generate document: {str(e)}",
            )

    async def _generate_txt(self, content: str, metadata: Dict) -> bytes:
        """Generate a plain text file."""
        lines = []

        # Add metadata header if present
        if metadata:
            lines.append("=" * 60)
            for key, value in metadata.items():
                lines.append(f"{key}: {value}")
            lines.append("=" * 60)
            lines.append("")

        lines.append(content)

        return "\n".join(lines).encode("utf-8")

    async def _generate_docx(self, content: str, title: str, metadata: Dict) -> bytes:
        """Generate a DOCX file."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX generation. "
                "Install with: pip install python-docx"
            )

        doc = Document()

        # Set document properties
        doc.core_properties.title = title
        if metadata.get("author"):
            doc.core_properties.author = metadata["author"]

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata if present
        if metadata:
            meta_para = doc.add_paragraph()
            for key, value in metadata.items():
                meta_para.add_run(f"{key}: {value}\n").italic = True

        doc.add_paragraph()  # Spacer

        # Parse content and add to document
        # Handle basic markdown-like formatting
        lines = content.split('\n')
        current_para = None

        for line in lines:
            stripped = line.strip()

            # Headers
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            # Horizontal rule
            elif stripped in ['---', '===', '***']:
                doc.add_paragraph('_' * 50)
            # Bullet points
            elif stripped.startswith('- ') or stripped.startswith('* '):
                doc.add_paragraph(stripped[2:], style='List Bullet')
            # Numbered list
            elif re.match(r'^\d+\.\s', stripped):
                text = re.sub(r'^\d+\.\s', '', stripped)
                doc.add_paragraph(text, style='List Number')
            # Empty line
            elif not stripped:
                if current_para:
                    current_para = None
                doc.add_paragraph()
            # Regular paragraph
            else:
                para = doc.add_paragraph()
                # Handle bold **text** and italic *text*
                self._add_formatted_text(para, stripped)

        # Save to bytes
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with basic markdown formatting to a paragraph."""
        # Simple regex-based formatting
        # Bold: **text**
        # Italic: *text*

        pattern = r'(\*\*.*?\*\*|\*.*?\*|[^*]+)'
        parts = re.findall(pattern, text)

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

    async def _generate_pdf(self, content: str, title: str, metadata: Dict) -> bytes:
        """Generate a PDF file."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
        except ImportError:
            # Fallback: convert text to simple PDF using basic method
            return await self._generate_simple_pdf(content, title, metadata)

        import io
        buffer = io.BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72,
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            alignment=TA_CENTER,
            spaceAfter=30,
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=12,
        )

        story = []

        # Title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))

        # Metadata
        if metadata:
            meta_text = "<br/>".join([f"<i>{k}: {v}</i>" for k, v in metadata.items()])
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 24))

        # Content - convert markdown-ish to reportlab
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()

            if stripped.startswith('### '):
                story.append(Paragraph(stripped[4:], styles['Heading3']))
            elif stripped.startswith('## '):
                story.append(Paragraph(stripped[3:], styles['Heading2']))
            elif stripped.startswith('# '):
                story.append(Paragraph(stripped[2:], styles['Heading1']))
            elif stripped.startswith('---') or stripped.startswith('==='):
                story.append(Spacer(1, 12))
                story.append(Paragraph('_' * 70, body_style))
                story.append(Spacer(1, 12))
            elif stripped.startswith('- ') or stripped.startswith('* '):
                story.append(Paragraph(f"• {stripped[2:]}", body_style))
            elif re.match(r'^\d+\.\s', stripped):
                story.append(Paragraph(stripped, body_style))
            elif not stripped:
                story.append(Spacer(1, 12))
            else:
                # Convert **bold** and *italic* to reportlab tags
                formatted = stripped
                formatted = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', formatted)
                formatted = re.sub(r'\*(.+?)\*', r'<i>\1</i>', formatted)
                story.append(Paragraph(formatted, body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    async def _generate_simple_pdf(self, content: str, title: str, metadata: Dict) -> bytes:
        """Fallback simple PDF generation without reportlab."""
        # If reportlab is not available, generate a text file
        # and return it (not ideal but functional)
        logger.warning("reportlab not installed, falling back to TXT format for PDF request")
        return await self._generate_txt(f"[PDF generation requires reportlab]\n\n{title}\n\n{content}", metadata)


class TemplateFillerTool(BaseTool):
    """
    Tool for filling document templates with data.

    Takes a template with {{placeholders}} and fills them with provided values.
    """

    @property
    def definition(self) -> ToolDefinition:
        """Return template filler tool definition."""
        return ToolDefinition(
            name="fill_template",
            description="""Fill a document template with data and generate a downloadable file.

Use this when you have a template with {{PLACEHOLDER}} markers that need to be filled with actual values.
The tool will replace all placeholders and generate the final document.""",
            parameters=[
                ToolParameter(
                    name="template_content",
                    type="string",
                    description="The template content with {{PLACEHOLDER}} markers",
                    required=True,
                ),
                ToolParameter(
                    name="values",
                    type="object",
                    description="Key-value pairs to fill in. Keys should match placeholder names (without braces)",
                    required=True,
                ),
                ToolParameter(
                    name="title",
                    type="string",
                    description="Document title/filename",
                    required=True,
                ),
                ToolParameter(
                    name="output_format",
                    type="string",
                    description="Output format: 'docx', 'pdf', or 'txt'",
                    required=False,
                    default="docx",
                    enum=["docx", "pdf", "txt"],
                ),
            ],
        )

    async def execute(self, **kwargs: Any) -> ToolResult:
        """Fill template and generate document."""
        template_content = kwargs.get("template_content", "")
        values = kwargs.get("values", {})
        title = kwargs.get("title", "document")
        output_format = kwargs.get("output_format", "docx")

        if not template_content:
            return ToolResult(
                success=False,
                output=None,
                error="Template content is required",
            )

        if not values:
            return ToolResult(
                success=False,
                output=None,
                error="Values dictionary is required to fill the template",
            )

        try:
            # Fill in placeholders
            filled_content = template_content
            unfilled_placeholders = []

            # Find all placeholders
            placeholders = re.findall(r'\{\{(\w+)\}\}', template_content)

            for placeholder in set(placeholders):
                if placeholder in values:
                    filled_content = filled_content.replace(
                        f"{{{{{placeholder}}}}}",
                        str(values[placeholder])
                    )
                else:
                    unfilled_placeholders.append(placeholder)

            # Warn about unfilled placeholders
            warnings = []
            if unfilled_placeholders:
                warnings.append(f"Unfilled placeholders: {', '.join(unfilled_placeholders)}")

            # Use DocumentGeneratorTool to create the file
            doc_tool = DocumentGeneratorTool()
            result = await doc_tool.execute(
                title=title,
                content=filled_content,
                output_format=output_format,
                metadata={"generated_from": "template", "filled_at": datetime.utcnow().isoformat()},
            )

            # Add warnings to output if any
            if result.success and warnings:
                result.output["warnings"] = warnings

            return result

        except Exception as e:
            logger.error("template_fill_failed", error=str(e))
            return ToolResult(
                success=False,
                output=None,
                error=f"Failed to fill template: {str(e)}",
            )
